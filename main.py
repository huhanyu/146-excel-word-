import math
import docx
import pandas as pd
from docx.enum.table import WD_ROW_HEIGHT, WD_ALIGN_VERTICAL
from docx.shared import Pt
from bakoverall import get_overall_table
from utils import set_cell_margins, set_cell_border
import openpyxl
####需要改的地方####
excel_path = 'CCWFO-03-21_12_14.xlsx'
save_filename='article_table1.docx'
# 读取Excel文件
# data = pd.read_excel("CCWFO-10-30_09_23.xlsx", sheet_name='overall')
data = pd.read_excel(excel_path, sheet_name='overall')
##################

# 获取算法个数
algorithms = data['Algrithm'].unique()
algorithm_count = len(algorithms)
print(algorithms)
# 获取函数个数
functions = data['F'].unique()
function_count = len(functions)
print(functions)
group_size = 3  # 每组的大小

num_groups = len(functions) // group_size  # 完整组的数量
remainder = len(functions) % group_size  # 剩余元素的数量

# 创建一个空的二维数组
two_dimensional_array = []

# 处理完整组
for i in range(num_groups):
    start_index = i * group_size
    end_index = (i + 1) * group_size
    group = functions[start_index:end_index]
    two_dimensional_array.append(group)
# 处理剩余元素
if remainder > 0:
    start_index = num_groups * group_size
    group = functions[start_index:]
    two_dimensional_array.append(group)

# 获取Avg,Std数据


n=math.ceil(function_count / 3) ##单位表格重复的次数 向上取整
unit_rows_num = algorithm_count+2 #单位行数量
unit_col_num = 9 #单位列数量

#创建word表格
doc = docx.Document()
# 添加一个n行9列的表格
table = doc.add_table(rows=unit_rows_num*n, cols=unit_col_num)

# 设置每行的第2、3、5、6、8、9列宽度为2.32厘米
columns_to_set_2_32_cm = [0,1, 2, 4, 5, 7, 8]  # 列索引从0开始计数
for row in table.rows:
    for col_idx in columns_to_set_2_32_cm:
        row.cells[col_idx].width = docx.shared.Cm(2.32)

# 设置每行的第4和第7列宽度为0.51厘米
columns_to_set_0_51_cm = [3, 6]  # 列索引从0开始计数
for row in table.rows:
    for col_idx in columns_to_set_0_51_cm:
        row.cells[col_idx].width = docx.shared.Cm(0.50)


# 设置每行的固定行高为0.5厘米
# fixed_row_height = docx.shared.Cm(0.5)
# for row in table.rows:
#     row.height = fixed_row_height



# 从第一行开始填入函数名称
print(len(functions))
for i, function_name in enumerate(functions):
    # 计算函数名称在表格中的位置
    row_index = i // 3 * unit_rows_num

    if (i + 1) % 3 == 1:col_index = 1
    elif (i + 1) % 3 == 2:col_index = 4
    elif (i + 1) % 3 == 0:col_index = 7  # 列索引为 1,4,7

    cell = table.cell(row_index, col_index)
    cell.text = function_name
    cell_avg = table.cell(row_index+1, col_index)
    cell_avg.text = 'Avg'
    cell_std = table.cell(row_index+1, col_index+1)
    cell_std.text = 'Std'
    ##设置表格框线
    cell_left = table.cell(row_index, col_index-1)
    cell_right = table.cell(row_index, col_index+1)
    set_cell_border(cell, top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                    bottom={"sz": 5, "color": "#000000", "val": "single"})
    set_cell_border(cell_left, top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                    bottom={"sz": 5, "color": "#000000", "val": "single"})
    set_cell_border(cell_right, top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                    bottom={"sz": 5, "color": "#000000", "val": "single"})
    if i == len(functions) - 1:
        for col_index1 in range (9):
            celll3= table.cell(row_index+3, col_index1)
            set_cell_border(celll3,bottom={"sz": 5, "color": "#000000", "val": "single"})


for i, algorithm_name in enumerate(algorithms):
    for j in range(n):
        row_index = i+2+j*unit_rows_num
        col_index = 0
        cell_alg=table.cell(row_index, col_index)
        cell_alg.text = algorithm_name

        for ii, function_name in enumerate(two_dimensional_array[j]): ##填入avg和std值

            mean_row_index = row_index
            std_row_index = mean_row_index

            if (ii + 1) % 3 == 1:
                mean_col_index = 1
            elif (ii + 1) % 3 == 2:
                mean_col_index = 4
            elif (ii + 1) % 3 == 0:
                mean_col_index = 7  # 列索引为 1,4,7

            std_col_index = mean_col_index+1

            f_alg_row = data[(data['F'] == function_name) & (data['Algrithm'] == algorithm_name)]
            mean_value = f_alg_row['mean'].values[0]
            mean_value = '%.4e' % mean_value
            std_value = f_alg_row['std'].values[0]
            std_value = '%.4e' % std_value

            cell_avg_value = table.cell(mean_row_index, mean_col_index)
            cell_avg_value.text = str(mean_value)
            cell_std_value = table.cell(std_row_index, std_col_index)
            cell_std_value.text = str(std_value)

# 遍历表格中的所有单元格

for row in table.rows:
    ##设置高度设置
    row.height_rule = WD_ROW_HEIGHT.EXACTLY
    row.height = docx.shared.Cm(0.51)
    for cell in row.cells:
        # 设置单元格中的文本对齐方式和垂直居中
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0, left=0, right=0)#设置单元格边距为0
        # 获取单元格的段落对象
        paragraph = cell.paragraphs[0]

        # 遍历段落中的所有运行对象
        for run in paragraph.runs:
            # 设置字体为Times New Roman
            run.font.name = 'Times New Roman'
            # 设置字号为五号
            run.font.size = Pt(10)

table2 = get_overall_table(excel_path,doc)

# table.style = 'Table Grid'
# 保存表格到 Word 文档
doc.save(save_filename)




print("算法个数：", algorithm_count)
print("函数个数：", function_count)












