import pandas as pd
import  math
import docx
from docx.shared import Pt
from docx.enum.table import WD_ROW_HEIGHT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from utils import set_cell_margins, set_cell_border

####需要改的地方####
filename='article_table.docx'
# 读取Excel文件
# data = pd.read_excel("CCWFO-10-30_09_23.xlsx", sheet_name='overall')
data = pd.read_excel("CCWFO-10-23_21_13.xlsx", sheet_name='overall')
data2 = pd.read_excel("CCWFO-10-23_21_13.xlsx", sheet_name='result & pValue')
excel_path = 'CCWFO-10-23_21_13.xlsx'
print(data2)
##################

def get_overall_table(excel_path,doc):

    data = pd.read_excel(excel_path, sheet_name='overall')
    data2 = pd.read_excel(excel_path, sheet_name='result & pValue')

    # 获取算法个数
    algorithms = data['Algrithm'].unique()
    algorithm_count = len(algorithms)
    # 获取函数个数
    functions = data['F'].unique()
    function_count = len(functions)


    # 创建一个新的 Word 文档
    # doc = docx.Document()

    # 添加一个带有 8 行 7 列的表格
    table = doc.add_table(rows=algorithm_count+2, cols=8)

    # 画表格
    columns_to_set_2_08_cm = [0,1,]  # 列索引从0开始计数
    for row in table.rows:
        for col_idx in columns_to_set_2_08_cm:
            row.cells[col_idx].width = docx.shared.Cm(2.08)
    columns_to_set_2_41_cm = [2]
    for row in table.rows:
        for col_idx in columns_to_set_2_41_cm:
            row.cells[col_idx].width = docx.shared.Cm(2.41)
    columns_to_set_2_09_cm = [3, 4, 6, 7]  # 列索引从0开始计数
    for row in table.rows:
        for col_idx in columns_to_set_2_09_cm:
            row.cells[col_idx].width = docx.shared.Cm(2.09)
    columns_to_set_0_32_cm = [5]
    for row in table.rows:
        for col_idx in columns_to_set_0_32_cm:
            row.cells[col_idx].width = docx.shared.Cm(0.32)


    # 填入默认字符并设置为粗体
    cell_1_2 = table.cell(0, 1)
    paragraph_1_2 = cell_1_2.paragraphs[0]
    run_1_2 = paragraph_1_2.add_run("Overall Rank")
    run_1_2.bold = True

    cell_2_2 = table.cell(1, 1)
    paragraph_2_2 = cell_2_2.paragraphs[0]
    run_2_2 = paragraph_2_2.add_run("RANK")
    run_2_2.bold = True

    cell_2_3 = table.cell(1, 2)
    paragraph_2_3 = cell_2_3.paragraphs[0]
    run_2_3 = paragraph_2_3.add_run("+/=-")
    run_2_3.bold = True

    cell_2_4 = table.cell(1, 3)
    paragraph_2_4 = cell_2_4.paragraphs[0]
    run_2_4 = paragraph_2_4.add_run("AVG")
    run_2_4.bold = True

    cells=table.rows[0].cells
    cells2=table.rows[algorithm_count+1].cells

    for cell in cells:
        set_cell_border(cell, top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},bottom={"sz": 5, "color": "#000000", "val": "single"})
    for cell2 in cells2:
        set_cell_border(cell2, bottom={"sz": 5, "color": "#000000", "val": "single"})


    for i, algorithm_name in enumerate(algorithms):
            row_index = i+2
            col_index = 0
            cell_alg=table.cell(row_index, col_index)
            cell_alg.text = algorithm_name
            ##写平均结果数据RANK
            avg_result_row = data2[(data2['F'] == '平均结果')]
            avg_result = avg_result_row[algorithm_name].values[0]
            avg_result = str(int(avg_result))
            cell_avg_result = table.cell(row_index, col_index+1)
            cell_avg_result.text = avg_result

            ##写+/=-数据abc
            cell_abc = table.cell(row_index, col_index + 2)
            if i ==0:
                cell_abc.text = '~'
            else:
                better_value = data2.iloc[data2[data2.iloc[:, 0] == 'CountOfBetter'].index[0], data2.columns.get_loc(algorithm_name) + 2]
                better_value = str(better_value)
                equel_value = data2.iloc[data2[data2.iloc[:, 0] == 'CountOfEqual'].index[0], data2.columns.get_loc(algorithm_name) + 2]
                equel_value = str(equel_value)
                worse_value = data2.iloc[data2[data2.iloc[:, 0] == 'CountOfWorse'].index[0], data2.columns.get_loc(algorithm_name) + 2]
                worse_value = str(worse_value)
                abc =  better_value + '/' + equel_value + '/' + worse_value
                cell_abc.text = abc

            ##写meanlevel数据AVG
            mean_level_row = data2[(data2['F'] == 'mean_level')]
            mean_level = mean_level_row[algorithm_name].values[0]
            mean_level = str(round(mean_level, 4))
            cell_mean_level = table.cell(row_index, col_index+3)
            cell_mean_level.text = mean_level

    # 遍历表格中的所有单元格，设置字体

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT.EXACTLY
        row.height = docx.shared.Cm(0.51)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=0, start=0, bottom=0, end=0, left=0, right=0)  # 设置单元格边距为0
            # 获取单元格的段落对象
            paragraph = cell.paragraphs[0]
            # 遍历段落中的所有运行对象
            for run in paragraph.runs:
                # 设置字体为Times New Roman
                run.font.name = 'Times New Roman'
                # 设置字号为五号
                run.font.size = Pt(10)

    return table
    # # 保存文档
    # doc.save("table.docx")